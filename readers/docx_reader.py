def read_docx(path):
    try:
        import docx2txt
        # docx2txt extracts paragraphs + textboxes
        txt = docx2txt.process(path)
        if txt and txt.strip():
            return txt
    except Exception:
        pass

    # fallback to python-docx (keep existing robust extraction)
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            for p in section.footer.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
    except Exception:
        pass
    return "\n".join(parts)